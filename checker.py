# -*- coding: utf-8 -*-
"""
EBYÜ Tez Formatlama Kontrolcüsü - Analiz Motoru (v3)

Kapsamlı kontroller:
- İçindekiler'den başlık tespiti
- Tablo/Şekil numaralandırma kontrolü
- Kaynakça format kontrolü
- Paragraf ve satır aralığı kontrolü
"""

import re
from typing import List, Dict, Any, Optional, Set, Tuple
from collections import defaultdict
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.shared import Pt, Cm
from docx.oxml.ns import qn

from config import ThesisConfig, FormatError, ErrorCategory, DEFAULT_CONFIG
from utils import (
    emu_to_cm, twips_to_cm, get_font_size_pt, get_text_snippet, count_words,
    is_chapter_heading, is_numbered_heading,
    is_table_caption, is_figure_caption, is_uppercase_text,
    is_dialogue_or_transcript, is_list_item, is_source_citation,
    has_visible_borders, StyleResolver
)


# Kapak sayfası desenleri
COVER_PATTERNS = [
    r"^T\.?C\.?$",
    r"ERZİNCAN.*ÜNİVERSİTESİ",
    r"SOSYAL BİLİMLER ENSTİTÜSÜ",
    r"ANA\s*BİLİM\s*DALI",
    r"BİLİM\s*DALI",
    r"YÜKSEK LİSANS",
    r"DOKTORA",
    r"TEZİ?$",
    r"HAZIRLAYAN",
    r"DANIŞMAN",
    r"^\d{4},?\s*ERZİNCAN$",
    r"^(OCAK|ŞUBAT|MART|NİSAN|MAYIS|HAZİRAN|TEMMUZ|AĞUSTOS|EYLÜL|EKİM|KASIM|ARALIK)\s+\d{4}",
]


class ThesisChecker:
    """EBYÜ Tez Formatlama Kontrolcüsü"""
    
    def __init__(self, config: ThesisConfig = None):
        self.config = config or DEFAULT_CONFIG
        self.errors: List[FormatError] = []
        self.document: Optional[Document] = None
        self.sections_found: Set[str] = set()
        self.abstract_text: str = ""
        self.cover_ended_at: int = 0
        self.total_checks: int = 0
        self.passed_checks: int = 0
        
        # İçindekiler'den çıkarılan başlıklar
        self.toc_headings: Dict[str, int] = {}  # başlık -> seviye
        self.tables_found: List[str] = []  # Tablo numaraları
        self.figures_found: List[str] = []  # Şekil numaraları
        self.in_references: bool = False  # Kaynakça bölümünde mi
        self.headings_found: List[str] = []  # Metinde bulunan başlıklar
        self.in_english_abstract: bool = False
        self.last_chapter_para_idx: int = -1
        self.resolver: Optional[StyleResolver] = None
        
        # Renk Şeması
        self.colors = {
            "LAYOUT": WD_COLOR_INDEX.YELLOW,      # Kenar boşluğu, girinti, aralık
            "STYLE": WD_COLOR_INDEX.RED,         # Yazı tipi, Boyut, Koyu/İtalik
            "REFERENCE": WD_COLOR_INDEX.TURQUOISE, # Kaynakça formatı
            "CONTENT": WD_COLOR_INDEX.BRIGHT_GREEN # Özet kelime sayısı vb.
        }
    
    def analyze(self, doc_path: str) -> Dict[str, Any]:
        """Tez dosyasını analiz eder."""
        self._reset()
        
        try:
            self.document = Document(doc_path)
            self.resolver = StyleResolver(self.document)
        except Exception as e:
            return self._error_report(str(e))
        
        # 1. Ön analiz
        self._find_cover_end()
        self._parse_toc()  # İçindekiler'i parse et
        self._find_sections()
        
        # 2. Kontroller
        self._check_abstract()
        self._check_margins()
        self._check_paragraphs()
        self._check_tables()
        self._check_table_figure_numbering()
        self._check_references()
        self._check_toc_consistency()
        self._check_page_numbers()
        self._check_footnotes()
        self._check_element_placement() # Yeni: Görsel yerleşim kontrolü
        
        return self._generate_report(), self.document
    
    def _reset(self):
        """Durumu sıfırla"""
        self.errors = []
        self.sections_found = set()
        self.abstract_text = ""
        self.cover_ended_at = 0
        self.total_checks = 0
        self.passed_checks = 0
        self.toc_headings = {}
        self.tables_found = []
        self.figures_found = []
        self.in_references = False
        self.headings_found = []
    
    def _find_cover_end(self):
        """Kapak sayfasının bitişini bul"""
        for i, para in enumerate(self.document.paragraphs):
            text = para.text.strip().upper()
            if any(x in text for x in ["BİLİMSEL ETİĞE", "ÖNSÖZ", "ÖN SÖZ", "ÖZET"]):
                self.cover_ended_at = i
                break
            if i > 50:
                self.cover_ended_at = 20
                break
    
    def _is_cover_or_skip(self, index: int, text: str) -> bool:
        """Paragrafın atlanması gerekip gerekmediğini kontrol et"""
        if index < self.cover_ended_at:
            return True
        text_upper = text.strip().upper()
        for pattern in COVER_PATTERNS:
            if re.search(pattern, text_upper, re.IGNORECASE):
                return True
        return False
    
    def _parse_toc(self):
        """İçindekiler'den başlıkları çıkar"""
        in_toc = False
        
        for i, para in enumerate(self.document.paragraphs):
            text = para.text.strip()
            text_upper = text.upper()
            
            # İçindekiler başlığını bul
            if "İÇİNDEKİLER" in text_upper and len(text) < 30:
                in_toc = True
                continue
            
            if in_toc:
                # Tablolar Listesi veya Şekiller Listesi gelince çık
                if any(x in text_upper for x in ["TABLOLAR LİSTESİ", "ŞEKİLLER LİSTESİ", "GİRİŞ"]) and len(text) < 30:
                    break
                
                # Numaralı başlıkları çıkar
                match = re.match(r'^(\d+(?:\.\d+)*\.?)\s+(.+?)(?:\s+\.+\s*\d+)?$', text)
                if match:
                    num = match.group(1)
                    title = match.group(2).strip()
                    level = num.count('.')
                    self.toc_headings[title.upper()] = level
    
    def _find_sections(self):
        """Gerekli bölümleri bul ve özeti çıkar"""
        required = {
            "ÖZET": "Özet",
            "ABSTRACT": "Abstract", 
            "GİRİŞ": "Giriş",
            "SONUÇ": "Sonuç",
            "KAYNAKÇA": "Kaynakça",
            "İÇİNDEKİLER": "İçindekiler",
        }
        
        in_abstract = False
        abstract_paragraphs = []
        ozet_found = False
        
        for para in self.document.paragraphs:
            text = para.text.strip()
            text_upper = text.upper()
            
            is_section_heading = False
            for key, name in required.items():
                if key in text_upper and len(text) < 50:
                    self.sections_found.add(name)
                    is_section_heading = True
                    
                    if key == "ÖZET" and "ABSTRACT" not in text_upper:
                        in_abstract = True
                        ozet_found = True
                    elif ozet_found and key != "ÖZET":
                        in_abstract = False
            
            if is_section_heading:
                continue
            
                # Anahtar Kelimeler gelince dur
                if "ANAHTAR KELİMELER" in text_upper:
                    in_abstract = False
                # Keywords veya ABSTRACT gelince dur
                elif any(x in text_upper for x in ["KEYWORDS", "ABSTRACT"]):
                    in_abstract = False
                # Tez başlığı gibi görünen büyük harfli uzun metin gelince dur
                elif len(text) > 100 and text.isupper() and not text.endswith("."):
                    in_abstract = False
                else:
                    abstract_paragraphs.append(text)
        
        self.abstract_text = " ".join(abstract_paragraphs)
        if "Abstract" in self.sections_found:
            self.in_english_abstract = True
    
    def _check_abstract(self):
        """Özet kelime sayısını kontrol et"""
        if not self.abstract_text:
            return
        
        word_count = count_words(self.abstract_text)
        self.total_checks += 1
        
        if word_count < self.config.abstract_min_words:
            self.errors.append(FormatError(
                category=ErrorCategory.ABSTRACT,
                message=f"Özet çok kısa: {word_count} kelime (minimum {self.config.abstract_min_words})",
                location="Özet",
                expected=f"En az {self.config.abstract_min_words} kelime",
                found=f"{word_count} kelime",
                snippet=get_text_snippet(self.abstract_text, 60)
            ))
        elif word_count > self.config.abstract_max_words:
            self.errors.append(FormatError(
                category=ErrorCategory.ABSTRACT,
                message=f"Özet çok uzun: {word_count} kelime (maksimum {self.config.abstract_max_words})",
                location="Özet",
                expected=f"En fazla {self.config.abstract_max_words} kelime",
                found=f"{word_count} kelime",
                snippet=get_text_snippet(self.abstract_text, 60)
            ))
        else:
            self.passed_checks += 1
    
    def _check_margins(self):
        """Kenar boşluklarını kontrol et"""
        if not self.document.sections:
            return
        
        section = self.document.sections[0]
        issues = []
        
        margins = [
            (section.top_margin, "Üst", self.config.margin_top),
            (section.bottom_margin, "Alt", self.config.margin_bottom),
            (section.left_margin, "Sol", self.config.margin_left),
            (section.right_margin, "Sağ", self.config.margin_right),
        ]
        
        for margin, name, expected in margins:
            self.total_checks += 1
            if margin is None:
                self.passed_checks += 1
                continue
            
            actual = margin.cm if hasattr(margin, 'cm') else emu_to_cm(margin)
            if actual and abs(actual - expected) > self.config.margin_tolerance_cm:
                issues.append(f"{name}: {actual:.1f} cm (olması gereken: {expected} cm)")
            else:
                self.passed_checks += 1
        
        if issues:
            self.errors.append(FormatError(
                category=ErrorCategory.MARGIN,
                message="; ".join(issues),
                location="Sayfa Düzeni",
                expected=f"{self.config.margin_top} cm",
                found="Farklı değerler",
                snippet="Sayfa kenar boşlukları ayarları"
            ))
    
    def _check_paragraphs(self):
        """Paragrafları kontrol et"""
        in_toc_list = False  # Tablolar/Şekiller Listesi içinde mi
        in_front_matter = True  # Ön sayfalar (önsöz, özet vs)
        
        for i, paragraph in enumerate(self.document.paragraphs):
            # Boş veya sadece whitespace içeren paragrafları atla
            text = paragraph.text.strip()
            if not text:
                continue
            
            # 2 karakterden kısa ve numerik olmayan metinleri atla (gürültü azaltma)
            if len(text) < 2 and not text.isdigit():
                continue
            
            text_upper = text.upper()
            
            # 1. Temel Font ve Boyut Kontrolü (Her zaman yap)
            # Kapak sayfaları ve başlıklar dahil tüm metin TNR 12pt (veya başlık kuralları) olmalı
            # Ancak skip kurallarından önce font ve boyutu en azından temel seviyede kontrol edelim
            
            # Gelişmiş: Skip öncesi sadece font kontrolü yapıp hataları biriktirebiliriz
            # Ama karmaşıklığı önlemek için önce skip kurallarını uygulayalım fakat 
            # Font hatası varsa skip etmemeyi tercih edebiliriz.
            
            # Alternatif: Font kontrolünü loop'un en başına alalım
            location = f"Paragraf {i + 1}"
            
            # Kapak ve özel sayfaları atla (diğer kurallar için)
            is_skipped_for_format = False
            if self._is_cover_or_skip(i, text):
                is_skipped_for_format = True
            
            # Tablolar/Şekiller Listesi bölümlerini atla
            if "TABLOLAR LİSTESİ" in text_upper or "ŞEKİLLER LİSTESİ" in text_upper:
                in_toc_list = True
                is_skipped_for_format = True
            if in_toc_list:
                if "GİRİŞ" in text_upper or "BÖLÜM" in text_upper:
                    in_toc_list = False
                else:
                    is_skipped_for_format = True
            
            # ÖNEMLİ: Font hatası skip kurallarına takılmamalı
            # Sadece çok kısa (1-2 harf) veya boş satırları font için atlayalım
            if len(text) > 3:
                f_err = self._check_font(paragraph)
                if f_err:
                    self._highlight_paragraph(paragraph, self.colors["STYLE"])
                    self.errors.append(FormatError(
                        category=f_err["category"],
                        message=f_err["message"],
                        location=location,
                        expected=f_err.get("expected", ""),
                        found=f_err.get("found", ""),
                        snippet=get_text_snippet(text, 80)
                    ))

            if is_skipped_for_format:
                continue
            
            # GİRİŞ bölümünden itibaren ana metin başlar
            if "GİRİŞ" in text_upper and len(text) < 20:
                in_front_matter = False
            
            if in_front_matter:
                alignment = paragraph.paragraph_format.alignment
                if alignment is not None and alignment.value in [1, 2]:  # CENTER veya RIGHT
                    continue
                if len(text) < 150:
                    continue
            
            # Kaynakça bölümünü işaretle
            if "KAYNAKÇA" in text_upper and len(text) < 20:
                self.in_references = True
                continue
            
            # === AKILLI FİLTRELER (Smart Filters) ===
            
            # 1. Diyalog / Transkript Atlaması
            if is_dialogue_or_transcript(text):
                # Diyaloglar için format kontrollerini -özellikle girinti ve aralık- atla
                self._highlight_paragraph(paragraph, WD_COLOR_INDEX.BRIGHT_GREEN)
                continue

            # 2. Kaynak Gösterimi Atlaması (Tablo/Şekil altı)
            if is_source_citation(text):
                self._highlight_paragraph(paragraph, WD_COLOR_INDEX.BRIGHT_GREEN)
                continue

            # 3. Liste Öğesi tespiti (Girinti kontrolü için kullanılacak)
            is_list = is_list_item(paragraph)

            para_issues = []
            
            # === 2. Başlık ve Özel Element Kontrolleri (Strict Segregation) ===
            is_cap = is_table_caption(text) or is_figure_caption(text)
            is_ch = is_chapter_heading(text) or (self.last_chapter_para_idx != -1 and i == self.last_chapter_para_idx + 1)
            is_num_h, _ = is_numbered_heading(text)
            
            if is_cap or is_ch or is_num_h:
                # Başlık-özel kontrolleri (Boyut, Font, Hizalama vb.)
                if is_table_caption(text):
                    self._check_table_caption(text, location)
                    para_issues.extend(self._check_caption_format(paragraph, text, "Tablo"))
                elif is_figure_caption(text):
                    self._check_figure_caption(text, location)
                    para_issues.extend(self._check_caption_format(paragraph, text, "Şekil"))
                elif is_ch:
                    self.headings_found.append(text.upper())
                    para_issues.extend(self._check_chapter_heading_format(paragraph, text, i))
                elif is_num_h:
                    self.headings_found.append(text.upper())
                    para_issues.extend(self._check_subheading_format(paragraph, text))

                # Hataları kaydet ve vurgula
                if para_issues:
                    self._highlight_paragraph(paragraph, self.colors["STYLE"])
                    for issue in para_issues:
                        self.errors.append(FormatError(
                            category=issue["category"],
                            message=issue["message"],
                            location=location,
                            expected=issue.get("expected", ""),
                            found=issue.get("found", ""),
                            snippet=get_text_snippet(text, 80)
                        ))
                
                # Bölüm başlığı sonrası özel kuralları (7cm, 4 satır) çalıştır
                if is_chapter_heading(text):
                    self._check_chapter_start_rules(i, location)
                
                # CRITICAL: Başlıklar paragraf kontrollerine girmemeli
                continue

            # === 3. Blok Alıntı Kontrolü ===
            if self._is_block_quote(paragraph):
                issues = self._check_block_quote_format(paragraph, text)
                if issues:
                    self._highlight_paragraph(paragraph, self.colors["LAYOUT"])
                    para_issues.extend(issues)
                # Blok alıntı da özel bir elementtir, normal paragraf kontrolünden geçebilir 
                # ama genellikle girinti ve aralıkları farklıdır.
                # Şimdilik blok alıntı sonrası da hataları kaydedip devam edelim.
                # Ancak kullanıcı normal paragraf kurallarının ona da uygulanmamasını isteyebilir.

            # === 4. Normal Paragraf Kontrolleri ===
            if not self.in_references:
                issues = self._check_normal_paragraph_format(paragraph, text)
                
                # Liste öğeleri için girinti ve satır aralığı hatasını filtrele
                if is_list:
                    issues = [iss for iss in issues if iss["category"] not in [ErrorCategory.PARAGRAPH, ErrorCategory.LINE_SPACING] or ("girinti" not in iss["message"].lower() and "aralık" not in iss["message"].lower())]
                    if not issues:
                        self._highlight_paragraph(paragraph, WD_COLOR_INDEX.BRIGHT_GREEN)
                
                if issues:
                    # Font hatası varsa KIRMIZI, yoksa SARI (Layout)
                    color = self.colors["STYLE"] if any("Yazı" in str(iss["category"]) for iss in issues) else self.colors["LAYOUT"]
                    self._highlight_paragraph(paragraph, color)
                para_issues.extend(issues)
            
            # Tüm hataları kaydet
            for issue in para_issues:
                # Çift kaydı önle: Başlıklar 'continue' ile bu kısma gelmez
                self.errors.append(FormatError(
                    category=issue["category"],
                    message=issue["message"],
                    location=location,
                    expected=issue.get("expected", ""),
                    found=issue.get("found", ""),
                    snippet=get_text_snippet(text, 80)
                ))

            # Epigraf Kontrolü
            from utils import is_epigraph
            if is_epigraph(paragraph):
                epigraph_issues = self._check_epigraph_format(paragraph, text)
                for issue in epigraph_issues:
                    self.errors.append(FormatError(
                        category=issue["category"],
                        message=issue["message"],
                        location=location,
                        expected=issue.get("expected", ""),
                        found=issue.get("found", ""),
                        snippet=get_text_snippet(text, 80)
                    ))
            
            # Kısa Alıntı Kontrolü (Tırnak kontrolü)
            from utils import is_short_quote
            if is_short_quote(text) and "“" not in text and "\"" not in text and len(text) > 20:
                # Basit bir kontrol: Eğer metin bir alıntı gibi duruyorsa ama tırnak yoksa uyar
                # (Bu çok gürültülü olabilir, şimdilik sadece plan dahilinde kalsın veya çok spesifik olsun)
                pass
    
    def _check_table_caption(self, text: str, location: str):
        """Tablo başlığı numaralandırmasını kontrol et"""
        # Tablo X.Y: veya Tablo X. Y: formatı
        match = re.match(r'^Tablo\s+(\d+)\.\s*(\d+)\s*:', text, re.IGNORECASE)
        if match:
            chapter = match.group(1)
            num = match.group(2)
            self.tables_found.append(f"{chapter}.{num}")
            self.passed_checks += 1
        else:
            self.total_checks += 1
            self.errors.append(FormatError(
                category=ErrorCategory.NUMBERING,
                message="Tablo numaralandırma formatı yanlış",
                location=location,
                expected="Tablo X.Y: Başlık",
                found=text[:40],
                snippet=get_text_snippet(text, 60)
            ))
    
    def _check_figure_caption(self, text: str, location: str):
        """Şekil başlığı numaralandırmasını kontrol et"""
        match = re.match(r'^Şekil\s+(\d+)\.\s*(\d+)\s*:', text, re.IGNORECASE)
        if match:
            chapter = match.group(1)
            num = match.group(2)
            self.figures_found.append(f"{chapter}.{num}")
            self.passed_checks += 1
        else:
            self.total_checks += 1
            self.errors.append(FormatError(
                category=ErrorCategory.NUMBERING,
                message="Şekil numaralandırma formatı yanlış",
                location=location,
                expected="Şekil X.Y: Başlık",
                found=text[:40],
                snippet=get_text_snippet(text, 60)
            ))
    
    def _check_table_figure_numbering(self):
        """Tablo ve şekil numaralandırma sırasını kontrol et"""
        # Her bölüm için sıralama kontrolü
        for item_type, items in [("Tablo", self.tables_found), ("Şekil", self.figures_found)]:
            if not items:
                continue
            
            by_chapter: Dict[str, List[int]] = defaultdict(list)
            for item in items:
                parts = item.split(".")
                if len(parts) == 2:
                    by_chapter[parts[0]].append(int(parts[1]))
            
            # Her bölümde sıralı mı kontrol et
            for chapter, nums in by_chapter.items():
                expected = list(range(1, len(nums) + 1))
                if sorted(nums) != expected:
                    self.errors.append(FormatError(
                        category=ErrorCategory.NUMBERING,
                        message=f"Bölüm {chapter}'de {item_type} numaralandırması sıralı değil",
                        location=f"{item_type} Numaralandırma",
                        expected=f"1, 2, 3, ...",
                        found=", ".join(map(str, sorted(nums))),
                        snippet=f"{item_type} dizisi: " + ", ".join(map(str, nums[:5]))
                    ))
    
    def _is_block_quote(self, para) -> bool:
        """Paragrafın bir blok alıntı olup olmadığını belirle (StyleResolver kullanarak)"""
        if not self.resolver:
            return False
            
        # EBYÜ: Her iki yandan 1.25cm girintili olan metinler blok alıntıdır
        left_indent = self.resolver.get_effective_paragraph_attribute(para, 'left_indent')
        right_indent = self.resolver.get_effective_paragraph_attribute(para, 'right_indent')
        
        left = left_indent.cm if hasattr(left_indent, 'cm') else (left_indent / 360000.0 if left_indent else 0.0)
        right = right_indent.cm if hasattr(right_indent, 'cm') else (right_indent / 360000.0 if right_indent else 0.0)
        
        # 1.0cm ile 1.5cm arası girinti varsa blok alıntı sayalım (Toleranslı 1.25cm)
        return left > 1.0 and right > 1.0

    def _check_block_quote_format(self, para, text: str) -> List[Dict]:
        """Blok alıntı formatını kontrol et (11pt, TNR, İtalik, 1.0 aralık, iki yana yaslı)"""
        issues = []
        
        # 1. Boyut (11pt)
        size_issue = self._check_size(para, self.config.font_size_block_quote, "Blok alıntı")
        if size_issue:
            issues.append(size_issue)
            
        # 2. İtalik kontrolü
        has_italic = all(run.font.italic for run in para.runs if run.text.strip())
        self.total_checks += 1
        if not has_italic:
            issues.append({
                "category": ErrorCategory.PARAGRAPH,
                "message": "Blok alıntı italik olmalıdır",
                "expected": "İtalik",
                "found": "Normal"
            })
        else:
            self.passed_checks += 1
            
        # 3. Satır Aralığı (1.0)
        self.total_checks += 1
        actual_ls = self.resolver.get_effective_line_spacing(para) if self.resolver else 1.0
        if abs(actual_ls - 1.0) > 0.15:
            issues.append({
                "category": ErrorCategory.PARAGRAPH,
                "message": "Blok alıntı satır aralığı 1.0 (tek) olmalı",
                "expected": "1.0",
                "found": f"{actual_ls:.1f}"
            })
        else:
            self.passed_checks += 1
            
        # 4. Hizalama (İki yana yaslı)
        alignment_issue = self._check_alignment(para, WD_ALIGN_PARAGRAPH.JUSTIFY, "Blok alıntı iki yana yaslı olmalı")
        if alignment_issue:
            issues.append(alignment_issue)
            
        return issues

    def _check_element_placement(self):
        """Görsel öğelerin (Tablo/Şekil) yerleşimini kontrol et - Sadece Tablo başlığı kontrolü"""
        if not self.document:
            return
            
        from docx.oxml.ns import nsmap
        from utils import is_table_caption, is_figure_caption
        
        body = self.document._element.body
        elements = body.xpath('.//w:p | .//w:tbl')
        
        table_count = 0
        for i, element in enumerate(elements):
            # TABLO KONTROLÜ (Başlık Üstte olmalı)
            if element.tag.endswith('tbl'):
                table_count += 1
                # Önceki 1-3 elemente bak (boşlukları atla)
                found_caption = False
                for j in range(1, 4):
                    if i - j >= 0:
                        prev = elements[i-j]
                        if prev.tag.endswith('p'):
                            text = "".join(prev.xpath('.//w:t/text()')).strip()
                            if not text: continue # Boş paragrafı atla
                            if is_table_caption(text):
                                found_caption = True
                                break
                            # Başka metin varsa dur (başlık yukarıda değil demek)
                            if len(text) > 20:
                                break
                
                if not found_caption:
                    self.total_checks += 1
                    self.errors.append(FormatError(
                        category=ErrorCategory.TABLE,
                        message="Tablo başlığı tablonun ÜSTÜNDE olmalıdır",
                        location=f"Tablo {table_count}",
                        expected="Üstte Başlık",
                        found="Eksik veya altta"
                    ))
        
        # ŞEKİL KONTROLÜ devre dışı - çok fazla false positive üretiyor
        # Word'de şekil yapısı çok değişken olabiliyor (inline, anchor, vb.)

    
    def _check_chapter_heading_format(self, para, text: str, para_idx: int = -1) -> List[Dict]:
        """Bölüm başlığı formatını kontrol et"""
        issues = []
        
        # Font kontrolü
        font_issue = self._check_font(para)
        if font_issue:
            issues.append(font_issue)
        
        # Boyut kontrolü - 14pt
        size_issue = self._check_size(para, self.config.font_size_chapter_heading, "Bölüm başlığı")
        if size_issue:
            issues.append(size_issue)
        
        # Koyu kontrolü
        self.total_checks += 1
        if not self._is_paragraph_bold(para):
            issues.append({
                "category": ErrorCategory.HEADING,
                "message": "Bölüm başlığı koyu olmalı",
                "expected": "Koyu",
                "found": "Normal"
            })
        else:
            self.passed_checks += 1
        
        # Bölüm başlığı (BÖLÜM X) sonrası gelen asıl başlık tespiti için işaretle
        from utils import is_chapter_title_only
        if is_chapter_title_only(text) and para_idx != -1:
            self.last_chapter_para_idx = para_idx
            
        # Büyük harf kontrolü
        self.total_checks += 1
        if not is_uppercase_text(text):
            issues.append({
                "category": ErrorCategory.HEADING,
                "message": "Bölüm başlığı tamamı büyük harf olmalı",
                "expected": "BÜYÜK HARF",
                "found": text[:30]
            })
        else:
            self.passed_checks += 1
        
        # Ortalama kontrolü
        alignment_issue = self._check_alignment(para, WD_ALIGN_PARAGRAPH.CENTER, "Bölüm başlığı ortalı olmalı")
        if alignment_issue:
            issues.append(alignment_issue)
        
        return issues
    
    def _check_subheading_format(self, para, text: str) -> List[Dict]:
        """Alt başlık formatını kontrol et"""
        issues = []
        
        font_issue = self._check_font(para)
        if font_issue:
            issues.append(font_issue)
        
        size_issue = self._check_size(para, self.config.font_size_subheading, "Alt başlık")
        if size_issue:
            issues.append(size_issue)
        
        self.total_checks += 1
        if not self._is_paragraph_bold(para):
            issues.append({
                "category": ErrorCategory.HEADING,
                "message": "Alt başlık koyu olmalı",
                "expected": "Koyu",
                "found": "Normal"
            })
        else:
            self.passed_checks += 1
        
        # Girinti ve Boşluk kontrolleri alt başlıklar için pas geçiliyor
        # (Kılavuza göre başlıkların kendine has kuralları var, gövde metni kuralları uygulanmaz)
        
        # Title Case kontrolü
        from utils import is_title_case
        if not is_title_case(text):
            self.total_checks += 1
            issues.append({
                "category": ErrorCategory.HEADING,
                "message": "Alt başlık her kelimesi büyük harfle başlamalı",
                "expected": "Her Kelime Büyük",
                "found": text[:30]
            })
        else:
            self.passed_checks += 1
            
        return issues
    
    def _check_caption_format(self, para, text: str, caption_type: str) -> List[Dict]:
        """Tablo/şekil başlığı formatını kontrol et"""
        issues = []
        
        font_issue = self._check_font(para)
        if font_issue:
            issues.append(font_issue)
        
        expected_size = self.config.font_size_table_caption if caption_type == "Tablo" else self.config.font_size_figure_caption
        size_issue = self._check_size(para, expected_size, f"{caption_type} başlığı")
        if size_issue:
            issues.append(size_issue)
        
        return issues
    
    def _check_normal_paragraph_format(self, para, text: str) -> List[Dict]:
        """Normal paragraf formatını kontrol et"""
        issues = []
        
        # Paragraf içi diğer kontroller (Genişlik, hizalama vs için font kontrolü zaten dış loopta yapıldı)
        # Ama yine de tutarlılık için burada font kontrolünü (dış loopta yapılmadıysa) tutabiliriz. 
        # Ancak dış loopta yapıldığı için burada sadece boyutu ve diğerlerini kontrol edelim.
        
        # Boyut kontrolü (12pt)
        size_issue = self._check_size(para, self.config.font_size_body, "Metin")
        if size_issue:
            issues.append(size_issue)
        
        # Hizalama kontrolü (iki yana yaslı)
        alignment_issue = self._check_alignment(para, WD_ALIGN_PARAGRAPH.JUSTIFY, "Metin iki yana yaslı olmalı")
        if alignment_issue:
            issues.append(alignment_issue)
        
        # Satır aralığı kontrolü (1.5)
        line_spacing_issue = self._check_line_spacing(para)
        if line_spacing_issue:
            issues.append(line_spacing_issue)
        
        # Paragraf girintisi kontrolü (1.25cm ilk satır)
        indent_issue = self._check_paragraph_indent(para)
        if indent_issue:
            issues.append(indent_issue)
        
        # Paragraf aralığı (nk) kontrolü (6nk önce / 6nk sonra)
        spacing_issue = self._check_paragraph_spacing(para)
        if spacing_issue:
            issues.append(spacing_issue)
        
        return issues
    
    def _check_paragraph_spacing(self, para) -> Optional[Dict]:
        """Paragraf boşluğu kontrolü - StyleResolver kullanarak kalıtımı çözer."""
        self.total_checks += 1
        
        if not self.resolver:
            return None

        # Effective attributes
        before_val = self.resolver.get_effective_paragraph_attribute(para, 'space_before')
        after_val = self.resolver.get_effective_paragraph_attribute(para, 'space_after')
        
        before = before_val.pt if hasattr(before_val, 'pt') else 0.0
        after = after_val.pt if hasattr(after_val, 'pt') else 0.0
        
        expected_before = self.config.paragraph_spacing_before
        expected_after = self.config.paragraph_spacing_after
        
        # Hata varsa (±1.1pt tolerans ile)
        if abs(before - expected_before) > 1.1 or abs(after - expected_after) > 1.1:
            return {
                "category": ErrorCategory.PARAGRAPH,
                "message": f"Paragraf aralığı {int(before)}nk-{int(after)}nk (olması gereken: {expected_before}nk-{expected_after}nk)",
                "expected": f"{expected_before}nk-{expected_after}nk",
                "found": f"{int(before)}nk-{int(after)}nk"
            }
        
        self.passed_checks += 1
        return None
    
    def _check_paragraph_indent(self, para) -> Optional[Dict]:
        """Paragraf ilk satır girintisi kontrolü - StyleResolver kullanarak kalıtımı çözer."""
        self.total_checks += 1
        
        if not self.resolver:
            return None

        first_line = self.resolver.get_effective_paragraph_attribute(para, 'first_line_indent')
        expected_cm = self.config.paragraph_first_line_indent  # 1.25
        
        if first_line is None:
            # Eğer hiç belirtilmemişse 0 sayılır, bu bir hatadır (EBYÜ'de girinti zorunlu)
            actual_cm = 0.0
        else:
            # cm'ye çevir
            if hasattr(first_line, 'cm'):
                actual_cm = first_line.cm
            else:
                actual_cm = first_line / 360000.0  # EMU to cm
        
        # Tolerans: 0.2cm
        if abs(actual_cm - expected_cm) > 0.2:
            return {
                "category": ErrorCategory.PARAGRAPH,
                "message": f"Paragraf girintisi {actual_cm:.2f}cm (olması gereken: {expected_cm}cm)",
                "expected": f"{expected_cm}cm",
                "found": f"{actual_cm:.2f}cm"
            }
        
        self.passed_checks += 1
        return None
    
    def _check_line_spacing(self, para) -> Optional[Dict]:
        """Satır aralığı kontrolü - StyleResolver kullanarak kalıtımı çözer."""
        self.total_checks += 1
        
        if not self.resolver:
            return None

        from utils import is_dialogue_or_transcript, is_source_citation, is_list_item
        text = para.text.strip()
        
        # === AKILLI ATLAMALAR (Human-like Bypass) ===
        if is_dialogue_or_transcript(text) or is_source_citation(text) or is_list_item(para):
            self.passed_checks += 1
            return None

        actual = self.resolver.get_effective_line_spacing(para)
        expected = self.config.line_spacing_body  # Varsayılan 1.5
        
        # Eğer dipnot veya tablo/şekil içindeyse expected 1.0 olabilir.
        # Bu fonksiyon şu an genel kullanılıyor, config'deki expected neyse ona bakıyor.
        # (Body metni için varsayılan 1.5 gelir)
        
        # Word'ün "1.5 satır aralığı" render farklılığına tolerans (1.25 - 1.65)
        if expected == 1.5:
            if 1.25 <= actual <= 1.65:
                self.passed_checks += 1
                return None
        # Word'ün "1.0 (tek) satır aralığı" toleransı (0.9 - 1.15)
        elif expected == 1.0:
            if 0.9 <= actual <= 1.15:
                self.passed_checks += 1
                return None
            
        if abs(actual - expected) > 0.15:
            return {
                "category": ErrorCategory.LINE_SPACING,
                "message": f"Satır aralığı {actual:.1f} (olması gereken: {expected})",
                "expected": f"{expected}",
                "found": f"{actual:.1f}"
            }
        
        self.passed_checks += 1
        return None
    
    def _check_references(self):
        """Kaynakça formatını kontrol et - APA 7 stili"""
        in_ref = False
        ref_count = 0
        ref_errors = []
        
        for i, para in enumerate(self.document.paragraphs):
            text = para.text.strip()
            
            if "KAYNAKÇA" in text.upper() and len(text) < 20:
                in_ref = True
                continue
            if in_ref and text:
                # Kaynakça bitti mi kontrol et (Yeni bölüm başlığı gelirse bitir)
                from utils import is_chapter_heading
                if is_chapter_heading(text):
                    break

                # 0. Gürültü Filtresi: Başlık ve boş satırları atla
                text_upper = text.upper()
                
                # Başlık tespiti: Genelde hepsi büyük ve kısa olur
                is_structural_header = (text_upper == text and len(text.split()) < 5)
                
                # Gürültü kelimeleri veya yapısal başlıklar
                noise_keywords = ["KISALTMALAR LİSTESİ", "ÖZ GEÇMİŞ", "ÖZGEÇMİŞ", "EK", "EKLER", "ÖNSÖZ", "ÖN SÖZ"]
                if any(x in text_upper for x in noise_keywords) and len(text) < 40:
                    continue
                
                # Kısa metinler, sadece numara olanlar veya kısaltma tanımları (A: B şeklinde)
                if is_structural_header or text.isdigit() or len(text.split()) < 3 or (":" in text and len(text.split(":")[0]) < 10):
                    continue

                ref_count += 1
                issues = []
                
                # 1. Asılı girinti kontrolü (1 cm)
                pf = para.paragraph_format
                if pf.first_line_indent:
                    indent_cm = pf.first_line_indent.cm if hasattr(pf.first_line_indent, 'cm') else pf.first_line_indent / 360000.0
                    # Asılı girinti için first_line_indent negatif olmalı
                    if indent_cm >= -0.1: # Negatif değilse asılı değildir
                        issues.append("Asılı girinti yok (1 cm asılı girinti olmalı)")
                    elif abs(abs(indent_cm) - 1.0) > 0.2:
                        issues.append(f"Asılı girinti {abs(indent_cm):.1f}cm (1 cm olmalı)")
                else:
                    issues.append("Asılı girinti eksik (1 cm asılı girinti olmalı)")
                
                # 2. İtalik kontrolü (APA: dergi/kitap adları italik olmalı)
                has_italic = False
                for run in para.runs:
                    if run.font.italic and run.text.strip():
                        has_italic = True
                        break
                
                if not has_italic:
                    issues.append("İtalik kısım yok (dergi/kitap adı italik olmalı)")
                
                # 3. Yıl formatı kontrolü - APA: (2021) veya (2021).
                import re
                if not re.search(r'\(\d{4}\)', text):
                    issues.append("Yıl formatı hatalı (YYYY) olmalı")
                
                # 4. Yazar formatı kontrolü - Soyad, A. veya Soyad, A.B. şeklinde başlamalı
                # Türkçe ve İngilizce soyadlar, tire içerebilir, çift baş harf olabilir
                author_pattern = r'^[A-ZÇĞİÖŞÜa-zçğıöşü][a-zçğıöşüA-ZÇĞİÖŞÜ\-\']+,\s*[A-ZÇĞİÖŞÜ]\.[A-ZÇĞİÖŞÜ]?\.?'
                if not re.match(author_pattern, text):
                    issues.append("Yazar formatı: Soyad, A. şeklinde olmalı")
                
                # 5. Boşluk kontrolü (3nk/3nk)
                pf = para.paragraph_format
                sb = pf.space_before.pt if pf.space_before else 0
                sa = pf.space_after.pt if pf.space_after else 0
                if abs(sb - 3) > 1 or abs(sa - 3) > 1:
                    issues.append(f"Kaynakça boşluğu {int(sb)}nk-{int(sa)}nk (3nk-3nk olmalı)")
                
                # TÜM kaynakları kontrol et, hataları kaydet
                if issues:
                    for issue in issues:
                        ref_errors.append({
                            "location": f"Kaynakça {ref_count}",
                            "message": issue,
                            "snippet": get_text_snippet(text, 50)
                        })
        
        # Hataları grupla ve ekle - Her hata için ayrı kayıt
        if ref_errors:
            # Hata türlerini ve snippet'leri grupla
            error_types = {}
            for err in ref_errors:
                msg = err["message"]
                if msg not in error_types:
                    error_types[msg] = []
                error_types[msg].append({
                    "loc": err["location"],
                    "snippet": err["snippet"]
                })
            
            for msg, refs in error_types.items():
                self.total_checks += 1
                # İlk 5 hatalı kaynağın snippet'ini göster
                snippets = [f"{r['loc']}: \"{r['snippet']}\"" for r in refs[:5]]
                more_text = f" (+{len(refs)-5} daha)" if len(refs) > 5 else ""
                self.errors.append(FormatError(
                        category=ErrorCategory.REFERENCE,
                        message=f"{msg} ({len(refs)} kaynak)",
                        location="Kaynakça",
                        expected="APA 7 formatı",
                        found=f"{len(refs)} kaynak",
                        snippet="; ".join(snippets) + more_text
                    ))
            
            # Kaynakça kısmını işaretle (Turkuaz)
            in_ref_section = False
            for para in self.document.paragraphs:
                if "KAYNAKÇA" in para.text.upper():
                    in_ref_section = True
                if in_ref_section:
                    if is_chapter_heading(para.text) and "KAYNAKÇA" not in para.text.upper():
                        break
                    if para.text.strip():
                        self._highlight_paragraph(para, self.colors["REFERENCE"])
    
    def _check_tables(self):
        """Tablo içeriklerini kontrol et - Akıllı Filtreleme ile"""
        from utils import has_visible_borders, is_table_caption
        
        for i, table in enumerate(self.document.tables):
            table_name = f"Tablo {i + 1}"
            
            # 1. Kenarlık ve Satır Sayısı Kontrolü
            # Görünür kenarlık yoksa veya 2 satırdan az ise kesinlikle LAYOUT tablodur
            if not has_visible_borders(table) or len(table.rows) < 2:
                continue
                
            # 2. Bağlam Kontrolü (Tablo metni içermeyen tabloları atla)
            has_table_marker = False
            for row in table.rows:
                for cell in row.cells:
                    if any(x in cell.text.upper() for x in ["TABLO", "TABLE"]):
                        has_table_marker = True
                        break
                if has_table_marker: break
            
            if not has_table_marker:
                continue

            # 3. Yerleşim Kontrolü
            
            wrong_sizes = set()
            wrong_fonts = set()
            
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if not para.text.strip():
                            continue
                        
                        for run in para.runs:
                            if not run.text.strip():
                                continue
                            
                            font = self.resolver.get_effective_font_name(run) if self.resolver else run.font.name
                            if font and font != self.config.font_name and font not in ["Symbol", "Wingdings"]:
                                wrong_fonts.add(font)
                            
                            size = self.resolver.get_effective_font_size(run) if self.resolver else (run.font.size.pt if run.font.size else None)
                            if size and abs(size - self.config.font_size_table_content) > 0.5:
                                wrong_sizes.add(int(size))
            
            issues = []
            if wrong_sizes:
                sizes_str = ", ".join([f"{s}pt" for s in sorted(wrong_sizes)])
                issues.append(f"Boyut: {sizes_str} (olması gereken: {self.config.font_size_table_content}pt)")
            
            if wrong_fonts:
                issues.append(f"Font: {', '.join(wrong_fonts)}")
            
            if issues:
                snippet = ""
                if table.rows:
                    snippet = get_text_snippet(table.rows[0].cells[0].text, 50)
                
                self.errors.append(FormatError(
                    category=ErrorCategory.TABLE,
                    message="; ".join(issues),
                    location=table_name,
                    expected=f"{self.config.font_size_table_content}pt, {self.config.font_name}",
                    found="Farklı format",
                    snippet=snippet
                ))
    
    def _is_paragraph_bold(self, para) -> bool:
        """Paragrafın koyu olup olmadığını kontrol et (StyleResolver kullanarak kalıtımla)"""
        if not self.resolver:
            return False

        # 1. Stil bizzat kalınlık dayatıyor mu? (e.g. Heading 1)
        if self.resolver.get_effective_font_bold(para) is True:
            return True

        # 2. Karakter bazlı eşik analizi
        bold_chars = 0
        total_chars = 0
        
        for run in para.runs:
            clean_text = run.text.strip()
            if not clean_text:
                continue
            
            text_len = len(clean_text)
            total_chars += text_len
            
            # Run bazlı kalınlık (Manuel + Karakter Stili + Paragraf Stili)
            if self.resolver.is_run_bold(run, para):
                bold_chars += text_len
        
        if total_chars == 0:
            return False
            
        # Paragrafın %80'inden fazlası koyu ise (dipnot/numara kaynaklı sapmaları önlemek için)
        return (bold_chars / total_chars) > 0.8
    
    def _check_font(self, para) -> Optional[Dict]:
        """Font kontrolü - StyleResolver kullanarak kalıtımı çözer."""
        self.total_checks += 1
        expected = self.config.font_name  # "Times New Roman"
        
        if not self.resolver:
            return None

        wrong_fonts = set()
        for run in para.runs:
            if not run.text.strip():
                continue
            
            font = self.resolver.get_effective_font_name(run)
            
            # Sembolleri ve özel durumları atla
            if font and font != expected and font not in ["Symbol", "Wingdings", "Cambria Math", "Webdings", "MS Mincho"]:
                wrong_fonts.add(font)
        
        if wrong_fonts:
            fonts_str = ", ".join(wrong_fonts)
            return {
                "category": ErrorCategory.FONT,
                "message": f"Yanlış yazı tipi: {fonts_str}",
                "expected": expected,
                "found": fonts_str
            }
        
        self.passed_checks += 1
        return None

    def _check_size(self, para, expected_pt: int, context: str) -> Optional[Dict]:
        """Boyut kontrolü - StyleResolver kullanarak kalıtımı çözer."""
        self.total_checks += 1
        tolerance = self.config.font_size_tolerance_pt  # 0.5pt
        
        if not self.resolver:
            return None

        wrong_sizes = set()
        for run in para.runs:
            if not run.text.strip():
                continue
            
            size = self.resolver.get_effective_font_size(run)
            
            if size and abs(size - expected_pt) > tolerance:
                wrong_sizes.add(size)
        
        if wrong_sizes:
            sizes_str = ", ".join([f"{s:.1f}pt" for s in sorted(wrong_sizes)])
            return {
                "category": ErrorCategory.FONT_SIZE,
                "message": f"{context} {sizes_str} (olması gereken: {expected_pt}pt)",
                "expected": f"{expected_pt}pt",
                "found": sizes_str
            }
        
        self.passed_checks += 1
        return None
    
    def _check_alignment(self, para, expected, message: str) -> Optional[Dict]:
        """Hizalama kontrolü"""
        self.total_checks += 1
        
        actual = para.paragraph_format.alignment
        
        if actual is None:
            self.passed_checks += 1
            return None
        
        if actual != expected:
            alignment_names = {
                WD_ALIGN_PARAGRAPH.LEFT: "Sola yaslı",
                WD_ALIGN_PARAGRAPH.CENTER: "Ortalı",
                WD_ALIGN_PARAGRAPH.RIGHT: "Sağa yaslı",
                WD_ALIGN_PARAGRAPH.JUSTIFY: "İki yana yaslı"
            }
            return {
                "category": ErrorCategory.PARAGRAPH,
                "message": message,
                "expected": alignment_names.get(expected, str(expected)),
                "found": alignment_names.get(actual, str(actual))
            }
        
        self.passed_checks += 1
        return None
    
    def _error_report(self, error_msg: str) -> Dict[str, Any]:
        """Hata raporu"""
        return {
            "total_errors": 1,
            "compliance_score": 0,
            "grouped_errors": {"Dosya Hatası": [{"location": "Dosya", "issues": [error_msg], "snippet": ""}]},
            "missing_sections": [],
            "abstract_issues": [],
            "sections_found": 0,
            "sections_required": 6
        }
    
    def _generate_report(self) -> Dict[str, Any]:
        """Rapor oluştur"""
        required_sections = ["Özet", "Abstract", "Giriş", "Sonuç", "Kaynakça", "İçindekiler"]
        missing = [s for s in required_sections if s not in self.sections_found]
        
        # Hataları grupla
        grouped: Dict[str, List[Dict]] = defaultdict(list)
        location_issues: Dict[str, Dict] = {}
        
        for error in self.errors:
            category = error.category.value
            loc_key = f"{category}_{error.location}"
            
            if loc_key in location_issues:
                if error.message not in location_issues[loc_key]["issues"]:
                    location_issues[loc_key]["issues"].append(error.message)
            else:
                location_issues[loc_key] = {
                    "category": category,
                    "location": error.location,
                    "issues": [error.message],
                    "snippet": error.snippet
                }
        
        for item in location_issues.values():
            grouped[item["category"]].append({
                "location": item["location"],
                "issues": item["issues"],
                "snippet": item["snippet"]
            })
        
        abstract_issues = [e.message for e in self.errors if e.category == ErrorCategory.ABSTRACT]
        
        if self.total_checks > 0:
            score = min(100.0, (self.passed_checks / self.total_checks) * 100)
        else:
            score = 100.0
        
        if missing:
            score = max(0, score - len(missing) * 5)
        
        return {
            "total_errors": len(self.errors),
            "total_checks": self.total_checks,
            "passed_checks": self.passed_checks,
            "compliance_score": round(score, 1),
            "missing_sections": missing,
            "sections_found": len(self.sections_found),
            "sections_required": len(required_sections),
            "abstract_issues": abstract_issues,
            "abstract_word_count": count_words(self.abstract_text) if self.abstract_text else 0,
            "tables_count": len(self.tables_found),
            "figures_count": len(self.figures_found),
            "toc_headings_count": len(self.toc_headings),
            "grouped_errors": dict(grouped)
        }

    def _check_toc_consistency(self):
        """İçindekiler listesinin metindeki başlıklarla tutarlılığını kontrol et"""
        if not self.toc_headings:
            return
            
        toc_list = list(self.toc_headings.keys())
        text_list = self.headings_found
        
        # Hızlı arama için temizlenmiş setler
        clean_text_set = {re.sub(r'[\.\s]', '', h) for h in text_list}
        clean_toc_set = {re.sub(r'[\.\s]', '', h) for h in toc_list}
        
        # İçindekiler'de olup metinde olmayanlar
        for toc_h in toc_list:
            clean_toc = re.sub(r'[\.\s]', '', toc_h)
            found = False
            if clean_toc in clean_text_set:
                found = True
            else:
                # Kısmi eşleşme kontrolü (daha yavaş ama güvenli)
                for ct in clean_text_set:
                    if clean_toc in ct or ct in clean_toc:
                        found = True
                        break
            
            self.total_checks += 1
            if not found:
                self.errors.append(FormatError(
                    category=ErrorCategory.SECTION,
                    message=f"İçindekiler'de yer alan başlık metinde bulunamadı: {toc_h}",
                    location="İçindekiler",
                    expected="Başlığın metinde yer alması",
                    found="Eksik başlık"
                ))
            else:
                self.passed_checks += 1
        
        # Metinde olup İçindekiler'de olmayanlar (Giriş sonrası)
        giriş_found = False
        for text_h in text_list:
            if "GİRİŞ" in text_h:
                giriş_found = True
            
            if giriş_found:
                clean_text = re.sub(r'[\.\s]', '', text_h)
                if clean_text in clean_toc_set:
                    self.passed_checks += 1
                    continue
                    
                found = False
                for ct in clean_toc_set:
                    if clean_text in ct or ct in clean_text:
                        found = True
                        break
                
                self.total_checks += 1
                if not found:
                    self.errors.append(FormatError(
                        category=ErrorCategory.SECTION,
                        message=f"Metindeki başlık İçindekiler'de bulunamadı: {text_h}",
                        location="İçindekiler",
                        expected="Başlığın İçindekiler'de yer alması",
                        found="Eksik İçindekiler kaydı"
                    ))
                else:
                    self.passed_checks += 1

    def _check_page_numbers(self):
        """Sayfa numaralarını kontrol et (10pt, TNR, Altta, Orta, 1.25cm)"""
        for i, section in enumerate(self.document.sections):
            self.total_checks += 1
            
            # 1. Mesafe kontrolü (1.25 cm)
            # Word ve yazıcı sürücüleri nedeniyle geniş bir tolerans (0.4cm - 2.5cm) uyguluyoruz.
            footer_dist = section.footer_distance.cm if section.footer_distance else 0
            if not (0.4 <= footer_dist <= 2.5):
                self.errors.append(FormatError(
                    category=ErrorCategory.NUMBERING,
                    message=f"Sayfa numarası (footer) mesafesi {footer_dist:.2f} cm (0.4 - 2.5 cm arası kabul edilir)",
                    location=f"Bölüm {i+1} Alt Bilgi",
                    expected="0.4 - 2.5 cm",
                    found=f"{footer_dist:.2f} cm"
                ))
            else:
                self.passed_checks += 1
            
            # 3. Numaralandırma Türü (i, ii vs. 1, 2)
            # Ön kısım (Küçük Roma), Giriş sonrası (Normal)
            # Not: python-docx ile start_type tespiti her zaman kolay değil, ama deniyoruz
            self.total_checks += 1
            # Basitçe ilk bölümler Roma, sonrakiler Arapça olmalı varsayımı (Tez yapısına göre)
            # Ancak dökümanda section ayrımı net olmayabilir.
            
            # 2. Footer içeriği (Stil ve hizalama)
            if section.footer and section.footer.paragraphs:
                p = section.footer.paragraphs[0]
                
                # Hizalama (CENTER olmalı)
                self.total_checks += 1
                if p.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                     self.errors.append(FormatError(
                        category=ErrorCategory.NUMBERING,
                        message="Sayfa numarası ortalanmış olmalı",
                        location=f"Bölüm {i+1} Sayfa Numarası",
                        expected="Ortalı",
                        found="Farklı hizalama"
                    ))
                else:
                    self.passed_checks += 1
                
                # Font (Times New Roman, 10pt)
                font_issue = self._check_font(p)
                if font_issue:
                    self.errors.append(FormatError(
                        category=ErrorCategory.FONT,
                        message="Sayfa numarası yazı tipi Times New Roman olmalı",
                        location=f"Bölüm {i+1} Sayfa Numarası",
                        expected="Times New Roman"
                    ))
                
                size_issue = self._check_size(p, 10, "Sayfa Numarası")
                if size_issue:
                    self.errors.append(FormatError(
                        category=ErrorCategory.FONT_SIZE,
                        message="Sayfa numarası 10 punto olmalı",
                        location=f"Bölüm {i+1} Sayfa Numarası",
                        expected="10 pt"
                    ))

    def _check_footnotes(self):
        """Dipnotları kontrol et (10pt, TNR, İki yana yaslı, 1.0 aralık, 0nk)"""
        try:
            # Footnotes part access
            footnotes_part = None
            for rel_id, part in self.document.part.related_parts.items():
                if "footnotes" in part.partname:
                    footnotes_part = part
                    break
            
            if not footnotes_part:
                return
            
            from docx.oxml import parse_xml
            from docx.oxml.ns import nsmap
            
            xml_content = footnotes_part.blob
            root = parse_xml(xml_content)
            # Sadece ID'si > 0 olan gerçek dipnotları kontrol et
            footnotes = root.xpath('//w:footnote[number(@w:id) > 0]', namespaces=nsmap)
            
            for footnote in footnotes:
                paragraphs = footnote.findall('.//w:p', nsmap)
                for p_xml in paragraphs:
                    # === GÜRÜLTÜ FİLTRESİ ===
                    # Metni bul (w:t elementlerini topla)
                    p_text = "".join(p_xml.xpath('.//w:t/text()', namespaces=nsmap)).strip()
                    # 1. Çok kısa metinler (Dipnot numarası veya separator çizgisi)
                    if not p_text or len(p_text) < 3:
                        continue
                    # 2. Sadece çizgiden oluşan veya boşluk içeren paragrafları atla
                    if all(char in '-_ \t\n\r' for char in p_text):
                        continue

                    # 1. Font ve Boyut (Runs)
                    runs = p_xml.findall('.//w:r', nsmap)
                    for r_xml in runs:
                        rPr = r_xml.find(qn('w:rPr'))
                        font = None
                        
                        if rPr is not None:
                            rf = rPr.find(qn('w:rFonts'))
                            if rf is not None:
                                font = rf.get(qn('w:ascii')) or rf.get(qn('w:hAnsi'))
                                if font is None:
                                    theme = rf.get(qn('w:asciiTheme')) or rf.get(qn('w:hAnsiTheme'))
                                    if theme:
                                        font = f"Tema Fontu ({theme})"
                            
                            # Boyut (sz 20 = 10pt)
                            sz = rPr.find(qn('w:sz'))
                            if sz is not None:
                                val = int(sz.get(qn('w:val'))) / 2
                                if val != 10:
                                    self._add_footnote_error("Dipnot yazı boyutu 10 punto olmalı", "10 pt", f"{val} pt")
                        
                        # Font hala bulunamadıysa döküman varsayılanlarına bak
                        if font is None:
                            try:
                                style_el = self.document.styles.element
                                rFonts_def = style_el.xpath('w:docDefaults/w:rPrDefault/w:rPr/w:rFonts')
                                if rFonts_def:
                                    def_f = rFonts_def[0].get(qn('w:ascii')) or rFonts_def[0].get(qn('w:asciiTheme'))
                                    if def_f and def_f != "Times New Roman":
                                        font = f"Varsayılan ({def_f})"
                            except Exception:
                                pass
                        
                        if font and font != "Times New Roman" and font not in ["Symbol", "Cambria Math"]:
                             self._add_footnote_error("Dipnot yazı tipi Times New Roman olmalı", "Times New Roman", font)
                    
                    pPr = p_xml.find(qn('w:pPr'))
                    
                    # Özellikleri stilden veya doğrudan pPr'den çek
                    def get_p_prop(p_xml, tag_name, attr_name=None):
                        pPr = p_xml.find(qn('w:pPr'))
                        if pPr is not None:
                            el = pPr.find(qn(tag_name))
                            if el is not None:
                                return el.get(qn(attr_name)) if attr_name else el
                        # Stilden dene
                        style_val = p_xml.find(qn('w:pPr/w:pStyle'))
                        if style_val is not None:
                            style_id = style_val.get(qn('w:val'))
                            # Basitlik için burada sadece doğrudan pPr'ye odaklanıyoruz 
                            # veya styles.xml'den çekebiliriz. Şimdilik pPr yeterli.
                        return None

                    if pPr is not None:
                        # Hizalama (jc) - 'both' veya 'distribute' (justify)
                        jc = pPr.find(qn('w:jc'))
                        val = jc.get(qn('w:val')) if jc is not None else 'left'
                        if val not in ['both', 'distribute']:
                            self._add_footnote_error("Dipnot iki yana yaslı olmalı", "İki yana yaslı", val)
                        
                        # Boşluk (spacing)
                        spacing = pPr.find(qn('w:spacing'))
                        if spacing is not None:
                            before = spacing.get(qn('w:before')) or '0'
                            after = spacing.get(qn('w:after')) or '0'
                            line = spacing.get(qn('w:line'))
                            rule = spacing.get(qn('w:lineRule'))

                            if before != '0' or after != '0':
                                self._add_footnote_error("Dipnot aralığı 0nk olmalı", "0nk-0nk", f"{before}nk-{after}nk")
                            
                            # Satır aralığı 1.0 (single) kontrolü
                            # rule=auto ve line=240, or rule=atLeast/exact and line=font_size
                            if rule == 'auto' and line and int(line) > 240:
                                self._add_footnote_error("Dipnot satır aralığı 1.0 (tek) olmalı", "1.0", f"{int(line)/240:.1f}")


        except Exception:
            pass

    def _add_footnote_error(self, msg, expected, found=""):
        self.total_checks += 1
        # Hata zaten listeye eklenmişse ekleme
        if any(e.message == msg for e in self.errors if e.location == "Dipnotlar"):
            return
            
        self.errors.append(FormatError(
            category=ErrorCategory.FOOTNOTE,
            message=msg,
            location="Dipnotlar",
            expected=expected,
            found=found
        ))
    def _check_spelling(self):
        """Yazım denetimi kaldırıldı (Zemberek bağımlılığı nedeniyle)."""
        pass

    def _check_chapter_start_rules(self, para_idx: int, location: str):
        """Bölüm başlığı öncesi 4 satır boşluk ve 7cm kuralını kontrol eder"""
        self.total_checks += 1
        paragraph = self.document.paragraphs[para_idx]
        
        # 1. 'Space Before' kontrolü (7cm ≈ 198pt - 3cm margin = 4cm ≈ 113pt ek boşluk)
        # 7cm kuralı: Sayfa başından itibaren 7cm. Margin 3cm ise, 4cm Space Before lazım.
        pf = paragraph.paragraph_format
        sb = pf.space_before.pt if pf.space_before else 0
        
        # Eğer sayfa başı tespiti yapamıyorsak (header/footer part tespiti zor), 
        # Space Before'un 100pt+ olmasını bekleyebiliriz.
        if sb < 100:
            # Alternatif: Önceki paragraflar boş mu?
            empty_count = 0
            for j in range(1, 6):
                if para_idx - j >= 0:
                    prev_text = self.document.paragraphs[para_idx - j].text.strip()
                    if not prev_text:
                        empty_count += 1
                    else:
                        break
            
            if empty_count < 4 and sb < 80:
                self.errors.append(FormatError(
                    category=ErrorCategory.MARGIN,
                    message="Bölüm başlığı öncesi 4 satır boşluk veya 7cm üst boşluk kuralı ihlali",
                    location=location,
                    expected="7cm üst boşluk / 4 satır boş",
                    found=f"{empty_count} boş satır, {int(sb)}nk boşluk"
                ))
            else:
                self.passed_checks += 1
        else:
            self.passed_checks += 1

    def _check_epigraph_format(self, para, text: str) -> List[Dict]:
        """Epigraf formatını kontrol et (11pt, İtalik, Sağa hizalı)"""
        issues = []
        
        # 1. Boyut (11pt)
        size_issue = self._check_size(para, 11, "Epigraf")
        if size_issue: issues.append(size_issue)
        
        # 2. İtalik
        self.total_checks += 1
        if not all(run.font.italic for run in para.runs if run.text.strip()):
            issues.append({
                "category": ErrorCategory.PARAGRAPH,
                "message": "Epigraf italik olmalı",
                "expected": "İtalik"
            })
        else:
            self.passed_checks += 1
            
        return issues


    def _highlight_paragraph(self, para, color):
        """Paragraftaki tüm run'ları belirtilen renkle vurgular."""
        for run in para.runs:
            run.font.highlight_color = color

    def _highlight_run(self, run, color):
        """Belirli bir run'ı vurgular."""
        run.font.highlight_color = color

def analyze_thesis(doc_path: str, config: ThesisConfig = None) -> Tuple[Dict[str, Any], Document]:
    """Tez dosyasını analiz et"""
    checker = ThesisChecker(config)
    return checker.analyze(doc_path)
