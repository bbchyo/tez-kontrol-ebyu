
import os
import sys
import json
from docx import Document
from docx.shared import Pt, Cm
from checker import analyze_thesis
from config import ThesisConfig

def run_test():
    doc_path = "/Users/halil/Desktop/tez kontrol/10516005 (1).docx"
    if not os.path.exists(doc_path):
        print(f"Error: {doc_path} not found.")
        return

    print("Phase 1: Analyzing original document...")
    config = ThesisConfig()
    results, marked_doc = analyze_thesis(doc_path, config)
    
    # Save initial report
    report_path = "/Users/halil/Desktop/tez kontrol/analiz_raporu_orijinal.txt"
    with open(report_path, "w", encoding="utf-8") as f:
        f.write("=== ORİJİNAL BELGE ANALİZİ ===\n")
        f.write(f"Compliance Score: {results['compliance_score']}%\n")
        f.write(f"Total Errors: {results['total_errors']}\n\n")
        
        for category, segments in results['grouped_errors'].items():
            f.write(f"[{category}]\n")
            for seg in segments:
                f.write(f"  Location: {seg['location']}\n")
                for issue in seg['issues']:
                    f.write(f"    - {issue}\n")
                if seg.get('snippet'):
                    f.write(f"    Snippet: {seg['snippet']}\n")
            f.write("\n")
            
    print(f"Original analysis saved to: {report_path}")

    print("Phase 2: Creating manipulated document...")
    doc = Document(doc_path)
    
    # 1. Heading manipulation: Remove bold from a numbered heading
    manipulated_count = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith("1.1.") or text.startswith("2.1."):
            for run in para.runs:
                run.font.bold = False
            manipulated_count += 1
            if manipulated_count >= 2: break
            
    # 2. Line spacing manipulation: Set very large spacing to a normal paragraph
    normal_p_count = 0
    for para in doc.paragraphs:
        if len(para.text) > 200:
            para.paragraph_format.line_spacing = 2.5
            normal_p_count += 1
            if normal_p_count >= 2: break
            
    # 3. Margin manipulation (on the fly is hard, but let's change a few font sizes)
    for para in doc.paragraphs:
        if "GİRİŞ" in para.text.upper() and len(para.text) < 20:
            for run in para.runs:
                run.font.size = Pt(16) # Should be 14
            break

    manipulated_doc_path = "/Users/halil/Desktop/tez kontrol/manipule_edilmis_tez.docx"
    doc.save(manipulated_doc_path)
    print(f"Manipulated document saved to: {manipulated_doc_path}")

    print("Phase 3: Analyzing manipulated document...")
    results_m, marked_doc_m = analyze_thesis(manipulated_doc_path, config)
    
    # Save manipulated report
    report_path_m = "/Users/halil/Desktop/tez kontrol/analiz_raporu_manipule.txt"
    with open(report_path_m, "w", encoding="utf-8") as f:
        f.write("=== MANİPÜLE EDİLMİŞ BELGE ANALİZİ ===\n")
        f.write(f"Compliance Score: {results_m['compliance_score']}%\n")
        f.write(f"Total Errors: {results_m['total_errors']}\n\n")
        
        for category, segments in results_m['grouped_errors'].items():
            f.write(f"[{category}]\n")
            for seg in segments:
                f.write(f"  Location: {seg['location']}\n")
                for issue in seg['issues']:
                    f.write(f"    - {issue}\n")
                if seg.get('snippet'):
                    f.write(f"    Snippet: {seg['snippet']}\n")
            f.write("\n")
            
    print(f"Manipulated analysis saved to: {report_path_m}")
    
    # Save marked doc
    marked_path = "/Users/halil/Desktop/tez kontrol/isaretlenmis_manipule_tez.docx"
    marked_doc_m.save(marked_path)
    print(f"Marked manipulated document saved to: {marked_path}")

if __name__ == "__main__":
    run_test()
