# -*- coding: utf-8 -*-
"""
EBYÜ Tez Formatlama Kontrolcüsü - AI İçerik Analizi Modülü

Google Gemini API kullanarak tez içeriğinin mantıksal analizini yapar.
"""

import re
from typing import Optional, Dict, List, Tuple
from docx import Document
from docx.shared import Pt, Twips

# Gemini API
try:
    import google.generativeai as genai
    GEMINI_AVAILABLE = True
except ImportError:
    GEMINI_AVAILABLE = False


class ThesisAIAnalyzer:
    """
    Gemini AI kullanarak tez içeriğini analiz eden sınıf.
    """
    
    # EBYÜ Kılavuzuna göre bölüm başlıkları
    SECTION_PATTERNS = {
        'ozet': r'^ÖZET$',
        'abstract': r'^ABSTRACT$',
        'giris': r'^GİRİŞ$|^BİRİNCİ BÖLÜM',
        'sonuc': r'^SONUÇ$|^SONUÇ VE ÖNERİLER$',
        'kaynakca': r'^KAYNAKÇA$|^KAYNAKLAR$',
    }
    
    def __init__(self, api_key: str):
        """
        Args:
            api_key: Google AI Studio API anahtarı
        """
        if not GEMINI_AVAILABLE:
            raise ImportError("google-generativeai paketi yüklü değil. 'pip install google-generativeai' ile yükleyin.")
        
        self.api_key = api_key
        genai.configure(api_key=api_key)
        
        # gemini-2.5-flash modeli - yeni model, ayrı quota
        self.model = genai.GenerativeModel('gemini-2.5-flash')
        self.doc = None
        self.full_text = ""
        self.sections: Dict[str, str] = {}
    
    def load_document(self, doc_path: str) -> None:
        """
        Word dökümanını yükler ve metni çıkarır.
        
        Args:
            doc_path: .docx dosya yolu
        """
        self.doc = Document(doc_path)
        self.full_text = self._extract_full_text()
        self.sections = self._extract_sections()
    
    def _extract_full_text(self) -> str:
        """Dökümanın tüm metnini çıkarır."""
        paragraphs = []
        for para in self.doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        return "\n\n".join(paragraphs)
    
    def _extract_sections(self) -> Dict[str, str]:
        """Önemli bölümleri ayıklar (Özet, Giriş, Sonuç vb.)"""
        sections = {}
        current_section = None
        current_content = []
        
        for para in self.doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Bölüm başlığı mı kontrol et
            found_section = None
            for section_name, pattern in self.SECTION_PATTERNS.items():
                if re.match(pattern, text, re.IGNORECASE):
                    found_section = section_name
                    break
            
            if found_section:
                # Önceki bölümü kaydet
                if current_section and current_content:
                    sections[current_section] = "\n".join(current_content)
                current_section = found_section
                current_content = []
            elif current_section:
                # Yeni bölüm başlığı görene kadar içeriği topla
                # Maksimum 50 paragraf
                if len(current_content) < 50:
                    current_content.append(text)
        
        # Son bölümü kaydet
        if current_section and current_content:
            sections[current_section] = "\n".join(current_content)
        
        return sections
    
    def check_abstract_page_overflow(self) -> Tuple[bool, str]:
        """
        Özet metninin tek sayfada olup olmadığını kontrol eder.
        
        Word'de gerçek sayfa düzeni render edilmeden kesin kontrolü yapmak zor,
        ancak tahmini satır/karakter sayısına göre değerlendirme yapılabilir.
        
        Returns:
            (is_overflow, message) - taşma durumu ve açıklama
        """
        ozet_text = self.sections.get('ozet', '')
        
        if not ozet_text:
            return False, "Özet bölümü bulunamadı."
        
        # Tahmini hesaplama:
        # - A4 sayfa, 3cm kenar boşlukları ile kullanılabilir alan: ~15cm x 24.7cm
        # - Times New Roman 12pt, 1.5 satır aralığı ile yaklaşık 28-30 satır/sayfa
        # - Ortalama satır uzunluğu: ~80 karakter
        
        lines = ozet_text.split('\n')
        total_chars = len(ozet_text)
        word_count = len(ozet_text.split())
        
        # Tahmini satır sayısı (wrapping dahil)
        estimated_lines = sum(max(1, len(line) // 80 + 1) for line in lines)
        
        # EBYÜ kuralı: Özet 200-250 kelime, tek sayfada olmalı
        MAX_LINES_PER_PAGE = 28  # 1.5 satır aralığı ile
        
        is_overflow = estimated_lines > MAX_LINES_PER_PAGE
        
        if is_overflow:
            message = (
                f"⚠️ Özet metni tek sayfayı aşıyor olabilir!\n"
                f"   - Tahmini satır sayısı: {estimated_lines} (max: {MAX_LINES_PER_PAGE})\n"
                f"   - Kelime sayısı: {word_count}\n"
                f"   - Karakter sayısı: {total_chars}"
            )
        else:
            message = (
                f"✅ Özet metni tek sayfa sınırları içinde görünüyor.\n"
                f"   - Tahmini satır sayısı: {estimated_lines}\n"
                f"   - Kelime sayısı: {word_count}"
            )
        
        return is_overflow, message
    
    def analyze_thesis_content(self) -> Dict:
        """
        Gemini AI kullanarak tez içeriğini analiz eder.
        
        Returns:
            Analiz sonuçlarını içeren sözlük
        """
        # Analiz için prompt hazırla
        prompt = self._build_analysis_prompt()
        
        try:
            response = self.model.generate_content(prompt)
            ai_analysis = response.text
        except Exception as e:
            ai_analysis = f"AI analizi sırasında hata oluştu: {str(e)}"
        
        # Özet sayfa kontrolü
        abstract_overflow, abstract_message = self.check_abstract_page_overflow()
        
        return {
            'ai_analysis': ai_analysis,
            'abstract_overflow': abstract_overflow,
            'abstract_message': abstract_message,
            'sections_found': list(self.sections.keys()),
            'total_words': len(self.full_text.split()),
            'total_chars': len(self.full_text),
        }
    
    def _build_analysis_prompt(self) -> str:
        """AI analizi için prompt oluşturur - EBYÜ 2022 Kılavuzu kurallarıyla."""
        
        # Bölümleri hazırla
        ozet = self.sections.get('ozet', 'Bulunamadı')[:2500]
        giris = self.sections.get('giris', 'Bulunamadı')[:4000]
        sonuc = self.sections.get('sonuc', 'Bulunamadı')[:4000]
        
        # Başlıkları ve şekil/tabloları ayıkla
        headings = self._extract_headings()
        figures_tables = self._extract_figures_tables()
        
        prompt = f"""Sen EBYÜ (Erzincan Binali Yıldırım Üniversitesi) Sosyal Bilimler Enstitüsü'nde görev yapan deneyimli bir tez danışmanısın. Aşağıdaki tezi EBYÜ 2022 Tez Yazım Kılavuzu'na göre detaylı şekilde değerlendir.

## EBYÜ 2022 TEZ YAZIM KILAVUZU KURALLARI

### ÖZET KURALLARI (Madde 2.1.8)
- Özet, tez hakkında "Ne, Niçin, Nasıl" sorularına cevap vermeli ve ulaşılan sonucu göstermeli
- Araştırma problemi (Ne), bu problemin hangi amaçla seçildiği (Niçin), izlenen yöntem (Nasıl), bulgular ve sonuçlara yer verilmeli
- Özet metni en az 200, en fazla 250 kelimeden oluşmalı ve BİR SAYFAYI AŞMAMALI
- En az 3, en fazla 5 anahtar kelime genelden özele sıralanmalı
- Tablo, şekil, grafik, formül, sembol kullanılmamalı

### BAŞLIK KURALLARI (Madde 1.5)
- Bölüm başlıkları (BİRİNCİ BÖLÜM, İKİNCİ BÖLÜM vb.) TAMAMI BÜYÜK HARF, 14 punto, koyu ve ortalı
- Bölüm başlıkları sayfa üst ilk satırından 4 satır boş bırakılarak 5. satırdan başlamalı
- Ana başlıklar ve alt başlıklar koyu, 1.25 cm girintili (paragraf başı ile hizalı), 12 punto
- Her kelimenin ilk harfi büyük olacak şekilde yazılmalı
- Bölüm numarasını içerecek şekilde numaralandırılmalı (örn: 2.3., 3.1.)

### TABLO VE ŞEKİL KURALLARI (Madde 1.7)
- Tablo ve şekiller sadece rakam kullanılarak koyu fontta numaralandırılmalı
- Her bölüm içinde kendi aralarında ayrı ayrı numaralandırılmalı (Tablo 1.1, Tablo 1.2, Şekil 2.1 vb.)
- Numara ile başlık arasında "iki nokta üst üste" işareti yer almalı (Tablo 1.1: Başlık)
- Tablo ve şekil başlıkları her kelimenin ilk harfi büyük
- Başlıklar Times New Roman 12 punto, içerik 11 punto
- Tablo ve şekiller tek sayfaya sığdırılmalı, sığmazsa "(devamı)" eklenmeli

### GİRİŞ KURALLARI (Madde 2.2.1)
- Tezin konusu, amacı, önemi ve yöntemi açık şekilde belirtilmeli
- Araştırma problemi soru cümlesi veya hipotez olarak net ifade edilmeli
- Giriş kısmında şekil, tablo vb. unsurlara yer verilmemeli
- "Giriş" başlığı dışında başlıklar ve alt başlıklar oluşturulmamalı

### SONUÇ KURALLARI (Madde 2.2.3)
- Elde edilen bulguların literatür ışığında genel yorumlanması yapılmalı
- Araştırma temel sorusu ve alt soruları esas alınarak cevaplar verilmeli
- Elde edilen sonuçlara göre önerilere yer verilebilir

### DOĞRUDAN ALINTI KURALLARI (Madde 1.9)
- 40 kelimeden az alıntılar çift tırnak içinde gösterilir
- 40 kelimeden fazla alıntılar blok alıntı olarak (her iki yandan 1.25 cm girintili, 11 punto, italik) düzenlenir

### AKADEMİK DİL VE ETİK (Madde 1.2, 1.10)
- Sade ve akıcı üslup, bilimsel dil kullanılmalı
- TDK Yazım Kılavuzu'na uyulmalı
- İntihal, sahtecilik, çarpıtma gibi etik ihlallerden kaçınılmalı

## ANALİZ EDİLECEK TEZ BÖLÜMLERİ

### ÖZET
{ozet}

### GİRİŞ
{giris}

### SONUÇ
{sonuc}

### TESPİT EDİLEN BAŞLIKLAR
{headings}

### TESPİT EDİLEN TABLO VE ŞEKİLLER
{figures_tables}

---

## DEĞERLENDİRME GÖREVİN

Yukarıdaki tez bölümlerini EBYÜ 2022 Kılavuzu kurallarına göre şu başlıklar altında değerlendir:

### 1. ÖZET DEĞERLENDİRMESİ
- Kelime sayısı 200-250 arasında mı?
- "Ne, Niçin, Nasıl" sorularına cevap veriyor mu?
- Araştırma problemi, amaç, yöntem, bulgular ve sonuç içeriyor mu?

### 2. BAŞLIK DEĞERLENDİRMESİ
- Bölüm başlıkları (BİRİNCİ BÖLÜM vb.) tamamen büyük harfle mi yazılmış?
- Alt başlıklar doğru numaralandırılmış mı (1.1., 1.2., 2.1. vb.)?
- Her kelimenin ilk harfi büyük mü?

### 3. TABLO VE ŞEKİL DEĞERLENDİRMESİ
- Numaralandırma formatı doğru mu (Tablo 1.1:, Şekil 2.1: vb.)?
- Bölüm bazında ayrı numaralandırma yapılmış mı?
- İki nokta üst üste kullanılmış mı?

### 4. GİRİŞ-SONUÇ TUTARLILIĞI
- Giriş'te belirtilen araştırma soruları Sonuç'ta yanıtlanmış mı?
- Mantıksal bütünlük sağlanmış mı?

### 5. AKADEMİK DİL VE ÜSLUP
- Bilimsel dil kullanılmış mı?
- Öznel ifadelerden kaçınılmış mı?

Her başlık için şu formatı kullan:
- ✅ UYGUN: (kısa açıklama)
- ⚠️ DİKKAT: (sorun ve öneri)
- ❌ SORUNLU: (detaylı açıklama ve düzeltme önerisi)

Değerlendirmeyi Türkçe olarak yap. Yapıcı ve somut öneriler sun.
"""
        return prompt
    
    def _extract_headings(self) -> str:
        """Tezdeki başlıkları tespit eder."""
        headings = []
        heading_patterns = [
            (r'^(BİRİNCİ|İKİNCİ|ÜÇÜNCÜ|DÖRDÜNCÜ|BEŞİNCİ)\s+BÖLÜM', 'Bölüm Başlığı'),
            (r'^(\d+\.)+\s*.+', 'Numaralı Başlık'),
            (r'^(GİRİŞ|SONUÇ|KAYNAKÇA|ÖZET|ABSTRACT)$', 'Ana Bölüm'),
        ]
        
        for para in self.doc.paragraphs:
            text = para.text.strip()
            if not text or len(text) > 200:
                continue
            
            for pattern, heading_type in heading_patterns:
                if re.match(pattern, text, re.IGNORECASE):
                    headings.append(f"[{heading_type}] {text}")
                    break
        
        if not headings:
            return "Başlık tespit edilemedi."
        
        return "\n".join(headings[:30])  # İlk 30 başlık
    
    def _extract_figures_tables(self) -> str:
        """Tezdeki tablo ve şekil başlıklarını tespit eder."""
        items = []
        patterns = [
            (r'^Tablo\s*(\d+\.?\d*)\s*[:\.]?\s*(.*)$', 'Tablo'),
            (r'^Şekil\s*(\d+\.?\d*)\s*[:\.]?\s*(.*)$', 'Şekil'),
            (r'^Grafik\s*(\d+\.?\d*)\s*[:\.]?\s*(.*)$', 'Grafik'),
            (r'^Resim\s*(\d+\.?\d*)\s*[:\.]?\s*(.*)$', 'Resim'),
        ]
        
        for para in self.doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            for pattern, item_type in patterns:
                match = re.match(pattern, text, re.IGNORECASE)
                if match:
                    num = match.group(1)
                    title = match.group(2) if len(match.groups()) > 1 else ''
                    items.append(f"[{item_type} {num}] {title}")
                    break
        
        if not items:
            return "Tablo veya şekil tespit edilemedi."
        
        return "\n".join(items[:30])  # İlk 30 öğe


def test_ai_analyzer(doc_path: str, api_key: str) -> Dict:
    """
    AI analizini test etmek için yardımcı fonksiyon.
    
    Args:
        doc_path: Test edilecek .docx dosya yolu
        api_key: Google AI API anahtarı
    
    Returns:
        Analiz sonuçları
    """
    analyzer = ThesisAIAnalyzer(api_key)
    analyzer.load_document(doc_path)
    results = analyzer.analyze_thesis_content()
    return results


if __name__ == "__main__":
    # Test için örnek kullanım
    import sys
    
    if len(sys.argv) < 3:
        print("Kullanım: python ai_analyzer.py <docx_path> <api_key>")
        sys.exit(1)
    
    doc_path = sys.argv[1]
    api_key = sys.argv[2]
    
    print(f"Dosya analiz ediliyor: {doc_path}")
    results = test_ai_analyzer(doc_path, api_key)
    
    print("\n" + "="*60)
    print("ÖZET SAYFA KONTROLÜ")
    print("="*60)
    print(results['abstract_message'])
    
    print("\n" + "="*60)
    print("AI ANALİZİ")
    print("="*60)
    print(results['ai_analysis'])
